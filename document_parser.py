import os
import pypdf
import docx
import pandas as pd

def chunk_text(text, chunk_size=1000, chunk_overlap=200):
    """
    Split text into overlapping chunks.
    """
    if not text:
        return []
    
    chunks = []
    start = 0
    text_len = len(text)
    
    while start < text_len:
        end = min(start + chunk_size, text_len)
        
        # If we aren't at the end of the text, try to find a natural boundary (newline or period)
        if end < text_len:
            # Look for last newline or period in the last 150 chars of the chunk
            search_area = text[end - 150:end]
            boundary = max(search_area.rfind('\n'), search_area.rfind('. '))
            if boundary != -1:
                end = end - 150 + boundary + 1
        
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
            
        start = end - chunk_overlap
        if start >= text_len - chunk_overlap:
            break
            
    return chunks

def parse_pdf(file_path):
    """
    Parse PDF file, returning a list of dictionaries with text and page numbers.
    """
    results = []
    try:
        reader = pypdf.PdfReader(file_path)
        for page_idx, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                # Store text by page for better citations
                results.append({
                    "text": text.strip(),
                    "page_num": page_idx + 1
                })
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {str(e)}")
    return results

def parse_docx(file_path):
    """
    Parse DOCX file, extracting paragraphs and tables.
    """
    try:
        doc = docx.Document(file_path)
        full_text = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
                
        # Extract tables and format them as pipe-delimited text
        for table in doc.tables:
            full_text.append("\n--- Table ---")
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_data:
                    full_text.append(" | ".join(row_data))
            full_text.append("-------------\n")
            
        combined_text = "\n".join(full_text)
        return [{"text": combined_text, "page_num": None}]
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}")

def parse_excel(file_path):
    """
    Parse Excel file sheets into Markdown-formatted tables.
    """
    try:
        xls = pd.ExcelFile(file_path)
        sheets_content = []
        
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name)
            if not df.empty:
                # Fill NaNs with empty string
                df = df.fillna("")
                # Convert sheet to string markdown table representation
                # We do a basic custom markdown table conversion to avoid dependencies if tabulate is missing
                headers = " | ".join(str(col) for col in df.columns)
                separator = " | ".join("---" for _ in df.columns)
                rows = []
                for _, row in df.iterrows():
                    rows.append(" | ".join(str(val) for val in row.values))
                
                sheet_md = f"Sheet Name: {sheet_name}\n" + headers + "\n" + separator + "\n" + "\n".join(rows)
                sheets_content.append(sheet_md)
                
        combined_text = "\n\n".join(sheets_content)
        return [{"text": combined_text, "page_num": None}]
    except Exception as e:
        raise ValueError(f"Failed to parse Excel file: {str(e)}")

def parse_csv(file_path):
    """
    Parse CSV file and format as table representation.
    """
    try:
        df = pd.read_csv(file_path)
        df = df.fillna("")
        
        headers = " | ".join(str(col) for col in df.columns)
        separator = " | ".join("---" for _ in df.columns)
        rows = []
        for _, row in df.iterrows():
            rows.append(" | ".join(str(val) for val in row.values))
            
        csv_md = headers + "\n" + separator + "\n" + "\n".join(rows)
        return [{"text": csv_md, "page_num": None}]
    except Exception as e:
        raise ValueError(f"Failed to parse CSV: {str(e)}")

def parse_text(file_path):
    """
    Read plain text file.
    """
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        return [{"text": content, "page_num": None}]
    except Exception as e:
        raise ValueError(f"Failed to parse text file: {str(e)}")

def parse_document(file_path):
    """
    Main entry point to parse a document based on its extension.
    Returns list of dicts: [{"text": str, "page_num": int or None}]
    """
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.pdf':
        return parse_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        return parse_docx(file_path)
    elif ext in ['.xlsx', '.xls']:
        return parse_excel(file_path)
    elif ext == '.csv':
        return parse_csv(file_path)
    elif ext in ['.txt', '.md', '.log', '.json', '.xml']:
        return parse_text(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")

def process_file_into_chunks(file_path, chunk_size=1000, chunk_overlap=200):
    """
    Parses the document and splits it into overlapping chunks, retaining metadata.
    Returns a list of dicts: [{"text": str, "page_num": int or None}]
    """
    parsed_sections = parse_document(file_path)
    all_chunks = []
    
    for section in parsed_sections:
        text = section["text"]
        page_num = section["page_num"]
        
        # Split section text into chunks
        chunks = chunk_text(text, chunk_size, chunk_overlap)
        for chunk in chunks:
            all_chunks.append({
                "text": chunk,
                "page_num": page_num
            })
            
    return all_chunks
